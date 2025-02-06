import streamlit as st
from transformers import pipeline
from PyPDF2 import PdfReader
from reportlab.pdfgen import canvas
import os
import base64
from docx import Document
import pandas as pd

# Function to set background
def set_background(image_path):
    with open(image_path, "rb") as image_file:
        encoded_string = base64.b64encode(image_file.read()).decode()
    st.markdown(
        f"""
        <style>
        header {{
            background-image: url("data:image/jpeg;base64,{encoded_string}"); 
            background-size: cover;
            background-repeat: no-repeat;
            background-attachment: fixed;
            height: 80px; 
        }}
        .stApp {{
            background-image: url("data:image/jpeg;base64,{encoded_string}"); 
            background-size: cover;
            background-repeat: no-repeat;
            background-attachment: fixed;
            font-family: 'Georgia', serif; /* Elegant default font */
        }}
        textarea {{
            font-family: 'Courier New', monospace; /* Set font for text areas */
        }}
        b {{
            font-weight: bold;
            font-family: 'Times New Roman', serif;
            color: #480607; 
            font-size: 24px;
        }}
        </style>
        """,
        unsafe_allow_html=True,
    )

# Set background
background_path = "pp.jpg"
if os.path.exists(background_path):
    set_background(background_path)
else:
    st.error("Background image not found. Please check the path.")

# Load redaction pipeline
st.title("Automated Redaction")
st.write("<b>Easily redact sensitive information from PDF, TXT, DOCX, or CSV files.</b>", unsafe_allow_html=True)

redactor = pipeline("token-classification", model="iiiorg/piiranha-v1-detect-personal-information")

label_mapping = {
    "I-ACCOUNTNUM": "Account Number",
    "I-BUILDINGNUM": "Building Number",
    "I-CITY": "City",
    "I-CREDITCARDNUMBER": "Credit Card Number",
    "I-DATEOFBIRTH": "Date of Birth",
    "I-DRIVERLICENSENUM": "Driver License Number",
    "I-EMAIL": "Email",
    "I-GIVENNAME": "Given Name",
    "I-IDCARDNUM": "ID Card Number",
    "I-PASSWORD": "Password",
    "I-SOCIALNUM": "Social Security Number",
    "I-STREET": "Street",
    "I-SURNAME": "Surname",
    "I-TAXNUM": "Tax Number",
    "I-TELEPHONENUM": "Telephone Number",
    "I-USERNAME": "Username",
    "I-ZIPCODE": "Zip Code",
    "O": "Other"
}

# Function to extract text from PDF
def extract_text_from_pdf(file_path):
    reader = PdfReader(file_path)
    text = ""
    for page in reader.pages:
        text += page.extract_text()
    return text

# Function to extract text from DOCX
def extract_text_from_docx(file_path):
    doc = Document(file_path)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text

# Function to extract text from TXT
def extract_text_from_txt(file_path):
    with open(file_path, "r") as file:
        return file.read()

# Function to extract text from CSV
def extract_text_from_csv(file_path):
    df = pd.read_csv(file_path)
    return df.to_csv(index=False)

# Function to redact sensitive information
def redact_and_group(text, selected_labels):
    detected_entities = redactor(text)
    detected_entities = sorted(detected_entities, key=lambda x: x['start'])

    redacted_spans = []
    current_span = None

    for entity in detected_entities:
        entity_start = entity['start']
        entity_end = entity['end']
        entity_label = entity.get('entity')

        if "all" in selected_labels or label_mapping.get(entity_label) in selected_labels:
            if current_span is None:
                current_span = {'start': entity_start, 'end': entity_end}
            elif entity_start <= current_span['end']:
                current_span['end'] = max(current_span['end'], entity_end)
            else:
                redacted_spans.append(current_span)
                current_span = {'start': entity_start, 'end': entity_end}

    if current_span is not None:
        redacted_spans.append(current_span)

    redacted_text = text
    for span in reversed(redacted_spans):
        redacted_length = span['end'] - span['start']
        redacted_text = redacted_text[:span['start']] + ("*" * redacted_length) + redacted_text[span['end']:]

    return redacted_text

# Function to create redacted PDF
def create_redacted_pdf(redacted_text, output_path):
    c = canvas.Canvas(output_path)
    y_position = 800
    for line in redacted_text.split("\n"):
        c.drawString(50, y_position, line)
        y_position -= 12
        if y_position < 50:
            c.showPage()
            y_position = 800
    c.save()

# Function to create redacted DOCX
def create_redacted_docx(redacted_text, output_path):
    doc = Document()
    for line in redacted_text.split("\n"):
        doc.add_paragraph(line)
    doc.save(output_path)

# Function to create redacted TXT
def create_redacted_txt(redacted_text, output_path):
    with open(output_path, "w") as file:
        file.write(redacted_text)

# Function to create redacted CSV
def create_redacted_csv(redacted_text, output_path):
    with open(output_path, "w") as file:
        file.write(redacted_text)

selected_labels = st.multiselect(
    "Select labels to redact:",
    options=["all"] + list(label_mapping.values()),
    default=["all"]
)

uploaded_file = st.file_uploader("Upload a file (PDF, TXT, DOCX, or CSV)", type=["pdf", "txt", "docx", "csv"])

if uploaded_file:
    file_extension = uploaded_file.name.split(".")[-1]
    temp_input_path = f"temp_input.{file_extension}"
    with open(temp_input_path, "wb") as f:
        f.write(uploaded_file.read())

    st.write("<b>Extracting text from the uploaded file...</b>", unsafe_allow_html=True)

    if file_extension == "pdf":
        extracted_text = extract_text_from_pdf(temp_input_path)
    elif file_extension == "txt":
        extracted_text = extract_text_from_txt(temp_input_path)
    elif file_extension == "docx":
        extracted_text = extract_text_from_docx(temp_input_path)
    elif file_extension == "csv":
        extracted_text = extract_text_from_csv(temp_input_path)

    st.text_area("Extracted Text", extracted_text, height=300)

    if st.button("Redact Text"):
        st.write("<b>Redacting sensitive information...</b>", unsafe_allow_html=True)
        redacted_text = redact_and_group(extracted_text, selected_labels)

        temp_output_path = f"redacted_output.{file_extension}"

        if file_extension == "pdf":
            create_redacted_pdf(redacted_text, temp_output_path)
        elif file_extension == "txt":
            create_redacted_txt(redacted_text, temp_output_path)
        elif file_extension == "docx":
            create_redacted_docx(redacted_text, temp_output_path)
        elif file_extension == "csv":
            create_redacted_csv(redacted_text, temp_output_path)

        with open(temp_output_path, "rb") as f:
            st.download_button(
                label="Download Redacted File",
                data=f,
                file_name=f"redacted_output.{file_extension}",
                mime=f"application/{file_extension}",
            )

        os.remove(temp_input_path)
        os.remove(temp_output_path)
